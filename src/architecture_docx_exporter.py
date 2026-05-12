"""Phase 4 (ADD): Export architecture design document to DOCX.

CLI:
    python src/architecture_docx_exporter.py [output_dir] [docx_path]
"""
import json
import os
import subprocess
import sys
import tempfile

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
sys.path.insert(0, SCRIPT_DIR)

from utils import os_type


# ── SVG → PNG via Puppeteer ───────────────────────────────────────────────────

_NODE_RENDER = """\
const puppeteer = require('puppeteer');
const fs = require('fs');
(async () => {
    const svg = fs.readFileSync(process.argv[2], 'utf8');
    const w = parseInt(svg.match(/width="(\\d+)"/)[1]);
    const h = parseInt(svg.match(/height="(\\d+)"/)[1]);
    const html = '<html><body style="margin:0;padding:0">' + svg + '</body></html>';
    const browser = await puppeteer.launch({args:['--no-sandbox','--disable-setuid-sandbox']});
    const page = await browser.newPage();
    await page.setViewport({width:w+1,height:h+1,deviceScaleFactor:2});
    await page.setContent(html,{waitUntil:'load'});
    await page.screenshot({path:process.argv[3],clip:{x:0,y:0,width:w,height:h}});
    await browser.close();
})();
"""


def _svg_to_png(svg_path: str) -> str | None:
    """Convert svg_path → png alongside it. Returns png path or None on failure."""
    png_path = os.path.splitext(svg_path)[0] + ".png"
    tmp = None
    try:
        # Write into project root so Node resolves node_modules/puppeteer
        tmp = os.path.join(PROJECT_ROOT, "_svg_to_png_tmp.js")
        with open(tmp, "w", encoding="utf-8") as f:
            f.write(_NODE_RENDER)
        cmd = ["node", tmp, svg_path, png_path]
        r = subprocess.run(
            cmd, capture_output=True, text=True, timeout=180,
            shell=(os_type == "Windows"), check=False,
            cwd=PROJECT_ROOT,
        )
        if r.returncode == 0 and os.path.isfile(png_path):
            return png_path
        err = (r.stderr or r.stdout or f"exit {r.returncode}").strip()
        print(f"[architecture_docx_exporter] warning: svg-to-png failed: {err}")
        return None
    except Exception as e:
        print(f"[architecture_docx_exporter] warning: svg-to-png error: {e}")
        return None
    finally:
        if tmp and os.path.isfile(tmp):
            try:
                os.unlink(tmp)
            except OSError:
                pass


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _add_para(doc, txt: str, style: str = "Normal"):
    return doc.add_paragraph(txt, style=style)


def _add_image(doc, png_path: str | None, inches: float = 5.0):
    from docx.shared import Inches
    if png_path and os.path.isfile(png_path):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(png_path, width=Inches(inches))
    else:
        _add_para(doc, "[Diagram not available]")


def _add_component_table(doc, groups: dict, font_pt):
    """Table: Component Group | Component | Description | Development Type | Note."""
    from docx.oxml.ns import qn
    headers = ("Component Group", "Component", "Description", "Development Type", "Note")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, label in enumerate(headers):
        hdr[i].text = label
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = font_pt

    # Track the first row index for each group (row 0 is the header)
    group_start: dict[str, int] = {}

    for group_name, components in groups.items():
        group_start[group_name] = len(table.rows)  # record before adding rows
        for i, (comp_name, comp_data) in enumerate(components.items()):
            description = comp_data.get("description", "") if isinstance(comp_data, dict) else ""
            row = table.add_row().cells
            row[0].text = group_name if i == 0 else ""
            row[1].text = comp_name
            row[2].text = description
            row[3].text = "New"
            row[4].text = ""
            for cell in row:
                for run in cell.paragraphs[0].runs:
                    run.font.size = font_pt

    # Merge Component Group cells vertically for groups with >1 component
    for group_name, components in groups.items():
        count = len(components)
        if count < 2:
            continue
        start_idx = group_start[group_name]
        end_idx = start_idx + count - 1
        top_cell = table.rows[start_idx].cells[0]
        bot_cell = table.rows[end_idx].cells[0]
        merged = top_cell.merge(bot_cell)
        # Center text vertically in merged cell
        tc = merged._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = tcPr.find(qn("w:vAlign"))
        if vAlign is None:
            from docx.oxml import OxmlElement
            vAlign = OxmlElement("w:vAlign")
            tcPr.append(vAlign)
        vAlign.set(qn("w:val"), "center")


# ── main export ───────────────────────────────────────────────────────────────

def export_architecture_docx(output_dir: str, docx_path: str) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("Error: python-docx not installed. pip install python-docx")
        raise SystemExit(1)

    static_dir = os.path.join(output_dir, "layer_static_diagrams")

    data_path = os.path.join(static_dir, "_layer_static_data.json")
    if not os.path.isfile(data_path):
        print(f"[architecture_docx_exporter] no data at {data_path}; run phase 3 first")
        raise SystemExit(1)

    with open(data_path, "r", encoding="utf-8") as f:
        layer_data = json.load(f)

    comp_design_path = os.path.join(static_dir, "_component_design_data.json")
    comp_design_data = {}
    if os.path.isfile(comp_design_path):
        with open(comp_design_path, "r", encoding="utf-8") as f:
            comp_design_data = json.load(f)

    font_pt = Pt(9)
    doc = Document()

    # ── title ─────────────────────────────────────────────────────────────────
    doc.add_heading("Software Architecture Design Specification", 0)

    # ── 1 Introduction ────────────────────────────────────────────────────────
    doc.add_heading("1 Introduction", level=1)
    doc.add_heading("1.1 Purpose", level=2)
    _add_para(doc, "[Purpose of this document.]")
    doc.add_heading("1.2 Scope", level=2)
    _add_para(doc, "[Scope of the software architecture design.]")
    doc.add_heading("1.3 Terms, Abbreviations and Definitions", level=2)
    _add_para(doc, "[Terms, abbreviations and definitions.]")

    # ── 3 Layer Design ────────────────────────────────────────────────────────
    doc.add_heading("3 Layer Design", level=1)

    for layer_idx, (layer_name, layer_info) in enumerate(layer_data.items(), start=1):
        svg_path = layer_info.get("svgPath")
        groups   = layer_info.get("groups", {})

        # 3.x LayerN
        doc.add_heading(f"3.{layer_idx} {layer_name}", level=2)

        # 3.x.1 Static Design
        doc.add_heading(f"3.{layer_idx}.1 Static Design", level=3)

        # render SVG → PNG then embed
        png_path = _svg_to_png(svg_path) if svg_path else None
        _add_image(doc, png_path)

        # component listing table
        if groups:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            hr = doc.add_paragraph()
            pPr = hr._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "auto")
            pBdr.append(bottom)
            pPr.append(pBdr)
            doc.add_heading("Component Information", level=4)
            _add_component_table(doc, groups, font_pt)

        # 3.x.3 Component Design
        layer_comp_designs = comp_design_data.get(layer_name, {})
        if layer_comp_designs:
            doc.add_heading(f"3.{layer_idx}.3 Component Design", level=3)
            comp_idx = 1
            for group_components in groups.values():
                for comp_name in group_components:
                    doc.add_heading(
                        f"3.{layer_idx}.3.{comp_idx} {comp_name}", level=4
                    )
                    svg_path = layer_comp_designs.get(comp_name, {}).get("svgPath")
                    png_path = _svg_to_png(svg_path) if svg_path else None
                    _add_image(doc, png_path, inches=5.5)
                    comp_idx += 1

    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)


def main():
    from core.paths import paths as _paths
    p = _paths()

    output_dir = os.path.join(p.output_dir, "add")
    docx_path  = os.path.join(p.output_dir, "Software Architecture Design Specification.docx")

    args = sys.argv[1:]
    if len(args) >= 1:
        output_dir = args[0]
    if len(args) >= 2:
        docx_path = args[1]

    if not os.path.isabs(output_dir):
        output_dir = os.path.join(PROJECT_ROOT, output_dir)
    if not os.path.isabs(docx_path):
        docx_path = os.path.join(PROJECT_ROOT, docx_path)

    export_architecture_docx(output_dir, docx_path)
    print(f"[architecture_docx_exporter] Written: {docx_path}")


if __name__ == "__main__":
    main()

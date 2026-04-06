"""Export interface_tables.json -> Software Detailed Design DOCX. Unit header table built from model."""
import os
import sys
import json
import subprocess
from typing import Optional, Tuple, List, Dict, Any

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "output")
MODEL_DIR = os.path.join(PROJECT_ROOT, "model")
COLS = ("Interface ID", "Interface Name", "Information", "Data Type", "Data Range", "Direction(In/Out)", "Source/Destination", "Interface Type")
# Placeholder when no value (no column may be empty)
NA = "N/A"


def _readable_label(name: str) -> str:
    """Convert an identifier like 'g_readWrite' or 'sb_index' into a human label."""
    if not name:
        return ""
    # Strip common global prefixes
    for prefix in ("g_", "s_", "t_"):
        if name.startswith(prefix):
            name = name[len(prefix) :]
            break
    # Replace underscores with spaces
    name = name.replace("_", " ")
    # Ignore very short/meaningless identifiers (e.g. "i", "v", "x")
    if len(name.strip()) <= 2:
        return ""
    # Basic title-case
    return name[:1].upper() + name[1:] if name else ""


def _strip_ext(path: str) -> str:
    if not path:
        return path
    return (path or "").replace("\\", "/")


def _path_no_ext(path: str) -> str:
    base, _ = os.path.splitext(path)
    return base.replace("\\", "/")


def _unit_paths(unit_info: dict) -> List[str]:
    """Path(s) without extension for this unit (for matching dataDictionary locations)."""
    path = unit_info.get("path")
    if not path:
        return []
    if isinstance(path, list):
        return [_strip_ext(p) for p in path]
    return [_strip_ext(path)]

def _struct_info_from_name(name: str) -> str:
    """Build a short description for a struct, e.g. 'HeapSort' -> 'Structure for Heap sorting'."""
    if not (name or "").strip():
        return "Structure for (unnamed)"
    s = name.strip()
    # Snake_case -> spaces; CamelCase -> spaces before capitals
    readable = []
    for i, c in enumerate(s):
        if c == "_":
            readable.append(" ")
        elif c.isupper() and i > 0 and readable and readable[-1] != " ":
            readable.append(" ")
            readable.append(c)
        else:
            readable.append(c)
    base = "".join(readable).strip()
    # Optional: lowercase trailing 'ing' context, e.g. "Heap Sort" -> "Heap sorting"
    if base and base.endswith(" Sort"):
        base = base[:-5] + " sorting"
    return f"Structure for {base}"


def _read_decl_snippet(abs_file: str, start_line: int, *, kind: str) -> str:
    """Extract declaration snippet safely using brace depth."""

    if not abs_file or not os.path.isfile(abs_file) or start_line < 1:
        return "-"

    try:
        with open(abs_file, "r", encoding="utf-8", errors="replace") as f:
            lines = f.readlines()
    except (OSError, IOError):
        return "-"

    if start_line > len(lines):
        return "-"

    i = start_line - 1
    buf = []
    max_lines = 60

    brace_depth = 0
    started = False

    for _ in range(max_lines):

        if i >= len(lines):
            break

        line = lines[i].rstrip("\n")
        stripped = line.strip()

        # skip leading comments
        if not started and (not stripped or stripped.startswith("//") or stripped.startswith("/*")):
            i += 1
            continue

        buf.append(line)

        if "{" in line:
            brace_depth += line.count("{")
            started = True

        if "}" in line:
            brace_depth -= line.count("}")

        # stop when typedef/struct block closes
        if started and brace_depth == 0 and ";" in line:
            break

        # stop simple declarations
        if not started and ";" in line:
            break

        i += 1

    out = "\n".join(buf).strip()

    # filter out function prototypes
    first_line = buf[0].strip() if buf else ""

    if "(" in first_line and ")" in first_line and first_line.endswith(";"):
        return "-"

    if kind == "typedef" and not out.lstrip().startswith("typedef"):
        return "-"

    if kind == "enum" and not (
        out.lstrip().startswith("enum") or out.lstrip().startswith("typedef enum")
    ):
        return "-"

    if kind == "struct":
        if not out.lstrip().startswith(("struct", "typedef struct")):
            return "-"

    return out if out else "-"

def _load_model_json(name: str) -> dict:
    path = os.path.join(MODEL_DIR, f"{name}.json")
    if not os.path.isfile(path):
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        return {}


def _load_base_path() -> str:
    meta = _load_model_json("metadata")
    return (meta.get("basePath") or "").strip()


def _load_abbreviations(project_root: str, config: dict) -> dict:
    """Load abbreviations from text file in config (llm.abbreviationsPath). Format: one per line, 'abbrev: meaning' or 'abbrev=meaning'; # = comment."""
    path = (config.get("llm") or {}).get("abbreviationsPath", "").strip()
    if not path:
        return {}
    full_path = os.path.join(project_root, path) if not os.path.isabs(path) else path
    if not os.path.isfile(full_path):
        return {}
    result = {}
    try:
        with open(full_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#"):
                    continue
                if ":" in line:
                    k, _, v = line.partition(":")
                elif "=" in line:
                    k, _, v = line.partition("=")
                else:
                    continue
                k, v = k.strip(), v.strip()
                if k:
                    result[k] = v
        return result
    except OSError:
        return {}


def _build_unit_header_table(
    unit_info: dict,
    interfaces: list,
    data_dictionary: dict,
    global_variables_data: dict,
    base_path: str,
    config: Optional[dict] = None,
    abbreviations: Optional[dict] = None,
) -> List[Dict[str, str]]:
    """Build rows for unit header table.

    - Column 1: full declaration as in code
    - Column 2: value (initializer / underlying type / enumerator values)
    """
    rows: List[Dict[str, str]] = []
    dd = data_dictionary or {}
    unit_paths_set = set(_unit_paths(unit_info))

    # Globals: use model/globalVariables.json so we can read exact line(s)
    for gid in unit_info.get("globalVariableIds", []) or []:
        g = (global_variables_data or {}).get(gid) or {}
        if (g.get("visibility") or "").lower() == "private":
            continue
        loc = g.get("location") or {}
        rel_file = (loc.get("file") or "").replace("\\", "/")
        line = int(loc.get("line") or 0)
        abs_file = os.path.join(base_path, rel_file) if base_path and rel_file else ""
        decl = _read_decl_snippet(abs_file, line, kind="var")
        info = g.get("value") or NA
        if (decl or "").strip() in ("", "-"):
            decl = g.get("qualifiedName") or g.get("name") or str(gid) or NA
        rows.append({"declaration": decl or NA, "information": info})

    # typedef, enum, define: match by unit file(s)
    for _type_name, t in dd.items():
        loc = t.get("location") or {}
        rel_file = (loc.get("file") or "").replace("\\", "/")
        type_file = _path_no_ext(rel_file)
        if not type_file or type_file not in unit_paths_set:
            continue
        kind = t.get("kind", "")
        # Include structs/unions so typedef-based structs (and unions) are visible
        # in the unit header table alongside typedef/enum/define entries.
        if kind not in ("typedef", "enum", "define"):
            continue
        line = int(loc.get("line") or 0)
        abs_file = os.path.join(base_path, rel_file) if base_path and rel_file else ""
        # Defines: use stored text/value from parser for exact macro
        if kind == "define":
            decl = t.get("text") or _read_decl_snippet(abs_file, line, kind="var")
            info = t.get("value", "") or NA
            if (decl or "").strip() in ("", "-"):
                decl = t.get("name") or _type_name or NA
            rows.append({"declaration": decl or NA, "information": info})
            continue

        decl = _read_decl_snippet(abs_file, line, kind=kind)

        if kind == "typedef":

            underlying = (t.get("underlyingType", "") or "").strip()

            # Only show values if typedef is aliasing an enum
            enum_ent = dd.get(underlying)

            if isinstance(enum_ent, dict) and enum_ent.get("kind") == "enum":
                enums = enum_ent.get("enumerators", []) or []
                parts = []
                for e in enums:
                    n = e.get("name", "")
                    v = e.get("value")
                    if n:
                        parts.append(f"{n}={v}" if v is not None else n)
                info = ", ".join(parts) if parts else NA

            elif isinstance(enum_ent, dict) and enum_ent.get("kind") == "struct":
                # typedef struct: description from name + fields (on the go, no store)
                type_name = t.get("name") or underlying or _type_name
                fields = enum_ent.get("fields") or []
                info = _struct_info_from_name(type_name)  # fallback
                if config:
                    try:
                        from llm_client import get_struct_description, _ollama_available
                        if _ollama_available(config) and config.get("llm", {}).get("descriptions", True):
                            llm_desc = get_struct_description(type_name, fields, config, abbreviations or {})
                            if llm_desc:
                                info = llm_desc
                    except ImportError:
                        pass
            else:
                info = NA
        elif kind == "enum":
            enums = t.get("enumerators", [])
            parts = []
            for e in enums:
                n = e.get("name", "")
                v = e.get("value")
                if n:
                    parts.append(f"{n}={v}" if v is not None else n)
            info = ", ".join(parts) if parts else NA
        else:
            info = NA

        if (decl or "").strip() in ("", "-"):
            decl = t.get("name") or _type_name or NA
        rows.append({"declaration": decl or NA, "information": info})

    # Deduplicate (same declaration can appear via enum + typedef entries)
    dedup = {}
    for r in rows:
        d = (r.get("declaration") or NA).strip()
        if d not in dedup:
            dedup[d] = r
        else:
            # Prefer the richer "name=value" info when both exist
            existing = dedup[d]
            existing_info = (existing.get("information") or "").strip()
            new_info = (r.get("information") or "").strip()
            if ("=" not in existing_info) and ("=" in new_info):
                dedup[d] = r
    out_rows = list(dedup.values())
    out_rows.sort(key=lambda r: (r.get("declaration") or "").lower())
    return out_rows


def _load_model_for_unit_headers() -> Tuple[dict, dict]:
    """Load units and dataDictionary from model/ for unit header table. Returns (units_data, data_dictionary)."""
    units_data = _load_model_json("units")
    data_dictionary = _load_model_json("dataDictionary")
    return units_data, data_dictionary


def _set_cell_font(cell, font_pt, bold=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = font_pt
            r.font.bold = bold


def _add_para(doc, text, style="Normal"):
    p = doc.add_paragraph(text, style=style)
    return p


def _add_mermaid_as_text(doc, mermaid: str, font_small):
    """Add Mermaid flowchart as monospace text block."""
    p = doc.add_paragraph()
    run = p.add_run(mermaid.strip())
    run.font.name = "Consolas"
    run.font.size = font_small


def _escape_mermaid_label_for_structure(text: str) -> str:
    t = (text or "").replace('"', "'").replace("\n", " ").replace("|", "\u00a6")
    return t


def _build_module_static_structure_mermaid(
    module_name: str,
    unit_rows: List[Tuple[Any, str, Any]],
) -> str:
    """Mermaid TB chart: one box for module, one row of child boxes for units.

    unit_rows: sorted list of (unit_key, unit_name_display, interfaces) per module.
    """
    mod_id = "MOD"
    mod_label = _escape_mermaid_label_for_structure(module_name)
    lines = [
        "%%{init: {'flowchart': {'ranksep': '0.55', 'nodesep': '0.35'}}}%%",
        "flowchart TB",
        f'  {mod_id}["{mod_label}"]',
    ]
    unit_ids = []
    for i, row in enumerate(unit_rows):
        disp = row[1] if len(row) > 1 else str(row[0])
        uid = f"U{i}"
        unit_ids.append(uid)
        lines.append(f'  {uid}["{_escape_mermaid_label_for_structure(disp)}"]')
        lines.append(f"  {mod_id} --> {uid}")
    lines.append(
        "  classDef moduleNode fill:#1e293b,stroke:#334155,color:#ffffff"
    )
    lines.append(
        "  classDef unitNode fill:#2563eb,stroke:#1d4ed8,color:#ffffff"
    )
    lines.append(f"  class {mod_id} moduleNode")
    if unit_ids:
        lines.append(f"  class {','.join(unit_ids)} unitNode")
    return "\n".join(lines)


def _parse_module_static_diagram_cfg(views_cfg: dict, export_cfg: dict = None) -> Tuple[bool, bool, float]:
    """enabled, renderPng, widthInches for Static Design module→units diagram (config.views.moduleStaticDiagram)."""
    raw = (views_cfg or {}).get("moduleStaticDiagram")
    if raw is None and export_cfg:
        raw = export_cfg.get("moduleStaticDiagram")
    if raw is None:
        raw = True
    if isinstance(raw, dict):
        return (
            bool(raw.get("enabled", True)),
            bool(raw.get("renderPng", True)),
            float(raw.get("widthInches", 5.5)),
        )
    return bool(raw), True, 5.5


def _render_mermaid_to_png(project_root: str, mermaid: str, png_path: str) -> bool:
    """Write .mmd and invoke mmdc; return True if png exists."""
    from utils import mmdc_path

    os.makedirs(os.path.dirname(png_path) or ".", exist_ok=True)
    mmd_path = os.path.splitext(png_path)[0] + ".mmd"
    try:
        with open(mmd_path, "w", encoding="utf-8") as f:
            f.write(mermaid)
    except OSError:
        return False
    puppeteer = os.path.join(project_root, "config", "puppeteer-config.json")
    mmdc = mmdc_path(project_root)
    cmd = [mmdc, "-i", mmd_path, "-o", png_path]
    if os.path.isfile(puppeteer):
        cmd.extend(["-p", puppeteer])
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=90, check=False)
        return r.returncode == 0 and os.path.isfile(png_path)
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False

def _add_flowchart_table(doc, func_name: str, description: str, input_name: str,
                         output_name: str, flowcharts: list, font_small):
    """Render a flowchart table matching the behaviour diagram table layout.

    Rows: Requirements (description + all flowcharts stacked), Capacity, Risk,
          Input Name, Output Name.

    flowcharts: list of (png_path_or_None, mermaid_str, label_str) tuples.
                First entry is the function's own flowchart; subsequent entries
                are private callee flowcharts.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
    from docx.shared import Inches

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"

    def _tight(para):
        para.paragraph_format.space_before = 0
        para.paragraph_format.space_after = 0
        para.paragraph_format.line_spacing = 1

    # Row 0: Requirements — description header then all flowcharts stacked
    row0 = table.rows[0].cells
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AUTO
    for cell in row0:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for para in cell.paragraphs:
            _tight(para)
    row0[0].text = "Requirements"
    _set_cell_font(row0[0], font_small)
    row0[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for para in row0[0].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _tight(para)

    # Right cell: description header, then each flowchart with an optional label
    p = row0[1].paragraphs[0]
    header_run = p.add_run(description.strip() if description and description.strip() else func_name or "-")
    header_run.bold = False
    header_run.font.size = font_small
    _tight(p)

    for png_path, mermaid, label in (flowcharts or []):
        if label:
            lp = row0[1].add_paragraph()
            label_run = lp.add_run(label)
            label_run.font.size = font_small
            _tight(lp)
        if png_path and os.path.isfile(png_path):
            ip = row0[1].add_paragraph()
            _tight(ip)
            try:
                ip.add_run().add_picture(png_path, width=Inches(4.0))
            except Exception:
                if mermaid:
                    fb = ip.add_run(mermaid.strip())
                    fb.font.name = "Consolas"
                    fb.font.size = font_small
        elif mermaid:
            mp = row0[1].add_paragraph()
            _tight(mp)
            fb = mp.add_run(mermaid.strip())
            fb.font.name = "Consolas"
            fb.font.size = font_small

    # Row 1: Risk
    row1 = table.rows[1].cells
    row1[0].text = "Risk"
    row1[1].text = "Medium"
    _set_cell_font(row1[0], font_small)
    _set_cell_font(row1[1], font_small)

    # Row 2: Capacity (Density)
    row2 = table.rows[2].cells
    row2[0].text = "Capacity(Density)"
    row2[1].text = "Common"
    _set_cell_font(row2[0], font_small)
    _set_cell_font(row2[1], font_small)

    # Row 3: Input Name
    row3 = table.rows[3].cells
    row3[0].text = "Input Name"
    row3[1].text = input_name or ""
    _set_cell_font(row3[0], font_small)
    _set_cell_font(row3[1], font_small)

    # Row 4: Output Name
    row4 = table.rows[4].cells
    row4[0].text = "Output Name"
    row4[1].text = output_name or ""
    _set_cell_font(row4[0], font_small)
    _set_cell_font(row4[1], font_small)


def _add_behavior_description_table(doc, behavior_description_list, input_name: str = "", output_name: str = ""):
    """Create a table with Requirements, Risk, Capacity, Input Name, Output Name rows.

    Args:
        doc: The python-docx Document object
        behavior_description_list: A list of strings to be displayed as bullet points
        input_name: Short human label for input of this behaviour
        output_name: Short human label for output of this behaviour
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"

    # Row 0: Requirements
    row0 = table.rows[0].cells
    row0[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    row0[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AUTO
    row0[0].text = "Requirements"
    for para in row0[0].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Remove extra spacing and fill second column for row 0
    for cell in row0:
        for para in cell.paragraphs:
            para.paragraph_format.space_before = 0
            para.paragraph_format.space_after = 0
            para.paragraph_format.line_spacing = 1

    # Second column: Bullet points from the list
    if behavior_description_list and isinstance(behavior_description_list, list):
        p = row0[1].paragraphs[0]
        header_run = p.add_run("Behavior Description\n")
        header_run.bold = True
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        for item in behavior_description_list:
            bullet_p = row0[1].add_paragraph()
            bullet_p.add_run(f"• {item}")
            bullet_p.paragraph_format.space_before = 0
            bullet_p.paragraph_format.space_after = 0
    else:
        row0[1].text = "-"

    # Row 1: Risk (default Medium for now)
    row1 = table.rows[1].cells
    row1[0].text = "Risk"
    row1[1].text = "Medium"

    # Row 2: Capacity (default Common for now)
    row2 = table.rows[2].cells
    row2[0].text = "Capacity"
    row2[1].text = "Common"

    # Row 3: Input Name (second column filled from model when available)
    row3 = table.rows[3].cells
    row3[0].text = "Input Name"
    row3[1].text = input_name or ""

    # Row 4: Output Name (second column filled from model when available)
    row4 = table.rows[4].cells
    row4[0].text = "Output Name"
    row4[1].text = output_name or ""

def _add_requirement_image_table(doc, png_path: str, flowchart_mermaid: str, font_small):
    import os
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
    from docx.shared import Inches

    width_inches = 4.0

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AUTO

    row = table.rows[0].cells

    # Vertical align TOP
    row[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    row[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    row[0].text = "Requirements"
    _set_cell_font(row[0], font_small)

    for para in row[0].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Remove extra spacing in both cells
    for cell in row:
        for para in cell.paragraphs:
            para.paragraph_format.space_before = 0
            para.paragraph_format.space_after = 0
            para.paragraph_format.line_spacing = 1

    # Clear second cell
    row[1].text = ""
    p = row[1].paragraphs[0]

    if png_path and os.path.isfile(png_path):
        try:
            run = p.add_run()
            run.add_picture(png_path, width=Inches(width_inches))
        except Exception:
            run = p.add_run(flowchart_mermaid.strip())
            run.font.name = "Consolas"
            run.font.size = font_small
    else:
        run = p.add_run(flowchart_mermaid.strip())
        run.font.name = "Consolas"
        run.font.size = font_small

def _load_flowcharts(flowcharts_dir: str) -> dict:
    """Return { unit_name: { func_name: flowchart_str } }."""
    result = {}
    if not os.path.isdir(flowcharts_dir):
        return result
    for fname in os.listdir(flowcharts_dir):
        if not fname.endswith(".json"):
            continue
        unit_name = fname[:-5]
        path = os.path.join(flowcharts_dir, fname)
        try:
            with open(path, "r", encoding="utf-8") as f:
                arr = json.load(f)
            if not isinstance(arr, list):
                continue
            result[unit_name] = {}
            for item in arr:
                name = (item.get("name") or "").strip()
                flowchart = (item.get("flowchart") or "").strip()
                if name and flowchart:
                    result[unit_name][name] = flowchart
        except (json.JSONDecodeError, OSError):
            pass
    return result


def _add_unit_header_table(doc, unit_header_rows: List[Dict[str, str]], font_small) -> None:
    """Add 2-column table under unit header: global variables/typedef/enum/define | information."""
    if not unit_header_rows:
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "global variables / typedef / enum / define"
    hdr[1].text = "information"
    _set_cell_font(hdr[0], font_small, bold=True)
    _set_cell_font(hdr[1], font_small, bold=True)
    for row_data in unit_header_rows:
        row = table.add_row().cells
        row[0].text = str(row_data.get("declaration") or NA)
        row[1].text = str(row_data.get("information") or NA)
        _set_cell_font(row[0], font_small)
        _set_cell_font(row[1], font_small)


def _add_interface_table(doc, interfaces, font_small):
    table = doc.add_table(rows=1, cols=len(COLS))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(COLS):
        hdr[i].text = c
        _set_cell_font(hdr[i], font_small, bold=True)

    for iface in interfaces:
        iface_type = iface.get("type", "") or "-"
        if "variableType" in iface:
            data_type = iface.get("variableType", "-") or "-"
            data_range = iface.get("range", "-") or "-"
        else:
            params = iface.get("parameters", [])
            data_type = "; ".join(p.get("type", "") for p in params) if params else "VOID"
            data_range = "; ".join(p.get("range", "") for p in params) if params else "NA"

        src_dest = iface.get("sourceDest") or "-"
        direction = iface.get("direction") or "-"
        info = iface.get("description", "") or "-"

        row = table.add_row().cells
        cells_text = (
            str(iface.get("interfaceId", "")),
            str(iface.get("interfaceName", "")),
            info,
            data_type,
            data_range,
            direction,
            src_dest,
            iface_type,
        )
        for i, txt in enumerate(cells_text):
            row[i].text = str(txt)
            _set_cell_font(row[i], font_small)


def _merge_vertical_cells(table, col: int, start_row: int, end_row: int) -> None:
    """Merge vertically from start_row to end_row inclusive (same column). Word keeps top-left cell content."""
    if start_row >= end_row:
        return
    top = table.cell(start_row, col)
    bottom = table.cell(end_row, col)
    top.merge(bottom)


def _add_component_unit_table(doc, component_name: str, unit_rows, font_small, config: dict, abbreviations: dict) -> None:
    """Add module-level table: Component | Unit | Description | Note.

    Description is derived from per-unit `interface_tables.json` entries:
    it aggregates available `description` fields from both functions and globals.
    Note column is left as N/A for now.
    """
    if not unit_rows:
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ("Component", "Unit", "Description", "Note")
    for i, h in enumerate(headers):
        hdr[i].text = h
        _set_cell_font(hdr[i], font_small, bold=True)

    def _cell_trim(s: str, max_len: int = 90) -> str:
        s = (s or "").strip()
        if not s:
            return ""
        if len(s) <= max_len:
            return s
        return s[: max_len - 3] + "..."

    def _unique_preserve_order(items: List[str]) -> List[str]:
        seen = set()
        out = []
        for it in items:
            if it in seen:
                continue
            seen.add(it)
            out.append(it)
        return out

    for row_idx, row_data in enumerate(unit_rows):
        # row_data = (unit_key, unit_name_display, interfaces)
        unit_name_display = row_data[1] if len(row_data) > 1 else str(row_data[0])
        interfaces = row_data[2] if len(row_data) > 2 else []

        # Pull descriptions from both functions and globals.
        fn_items: List[Tuple[str, str]] = []
        gv_items: List[Tuple[str, str]] = []
        for iface in interfaces or []:
            d = str(iface.get("description") or "").strip()
            if not d or d in ("-", NA):
                continue

            # Clean up newlines so prompt/table content stays readable.
            d_clean = " ".join(d.split())

            iface_type = iface.get("type") or ""
            name = (
                iface.get("interfaceName")
                or iface.get("name")
                or (iface.get("qualifiedName") or "").split("::")[-1]
                or ""
            )
            name = str(name).strip()

            if iface_type == "Global Variable":
                gv_items.append((name, d_clean))
            else:
                # Default to "Function" for any non-global interface types we got.
                fn_items.append((name, d_clean))

        # De-duplicate while preserving order.
        def _dedup_items(items: List[Tuple[str, str]]) -> List[Tuple[str, str]]:
            seen = set()
            out = []
            for n, d in items:
                key = (n or "").strip().lower() + "|" + (d or "").strip().lower()
                if key in seen:
                    continue
                seen.add(key)
                out.append((n, d))
            return out

        fn_items = _dedup_items(fn_items)
        gv_items = _dedup_items(gv_items)

        description_text = NA
        try:
            from llm_client import _ollama_available, get_unit_description

            if _ollama_available(config):
                description_text = get_unit_description(
                    unit_name_display,
                    fn_items,
                    gv_items,
                    config,
                    abbreviations or {},
                ).strip()
        except Exception:
            # Fall back to deterministic join if the LLM call fails.
            description_text = NA

        if not description_text or description_text in ("-", NA):
            # Fallback: join all descriptions (not just 3), but cap length for DOCX readability.
            all_descs = [d for _, d in (fn_items + gv_items)]
            all_descs = _unique_preserve_order(all_descs)
            joined = "; ".join(all_descs)
            description_text = _cell_trim(joined, max_len=120)

        # Keep output short even when the LLM returns long text.
        description_text = _cell_trim(str(description_text or NA), max_len=140)
        if not description_text:
            description_text = NA

        note_text = NA

        row = table.add_row().cells
        # Component is one per module; only first body row holds text, then column 0 is merged.
        row[0].text = str(component_name or NA) if row_idx == 0 else ""
        row[1].text = str(unit_name_display or NA)
        row[2].text = str(description_text or NA)
        row[3].text = str(note_text or NA)
        for c in row:
            _set_cell_font(c, font_small)

    # Merge Component column across all body rows (same component for the whole table).
    n = len(table.rows)
    if n > 2:
        _merge_vertical_cells(table, 0, 1, n - 1)


def export_docx(json_path: str = None, docx_path: str = None, selected_group: str | None = None) -> Tuple[bool, Optional[str]]:
    from utils import load_config, safe_filename, KEY_SEP
    config = load_config(PROJECT_ROOT)
    export_cfg = config.get("export", {})
    json_path = json_path or os.path.join(OUTPUT_DIR, "interface_tables.json")
    json_path = os.path.abspath(json_path)
    # Views write next to interface_tables.json (e.g. output/<group>/); do not use output/ only.
    artifacts_dir = os.path.dirname(json_path)
    # Allow group-specific filenames via {group} placeholder.
    group_name = selected_group or "all"
    raw_docx = export_cfg.get("docxPath", "output/software_detailed_design.docx")
    raw_docx = raw_docx.replace("{group}", group_name)
    if not docx_path:
        docx_path = os.path.join(PROJECT_ROOT, raw_docx)
    font_size = int(export_cfg.get("docxFontSize", 8))
    views_cfg = config.get("views", {})
    msd_enabled, msd_render_png, msd_width_in = _parse_module_static_diagram_cfg(views_cfg, export_cfg)
    fc_cfg = views_cfg.get("flowcharts") if isinstance(views_cfg.get("flowcharts"), dict) else {}
    flowcharts_enabled = bool(views_cfg.get("flowcharts"))
    flowcharts_render_png = fc_cfg.get("renderPng", True)
    flowcharts_dir = os.path.abspath(os.path.join(artifacts_dir, "flowcharts"))
    flowcharts_map = _load_flowcharts(flowcharts_dir) if flowcharts_enabled else {}

    if not os.path.isfile(json_path):
        print(f"Error: {json_path} not found. Run pipeline first.")
        return (False, None)

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        font_small = Pt(font_size)
    except ImportError:
        print("Error: python-docx not installed. pip install python-docx")
        return (False, None)

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    abbreviations = _load_abbreviations(PROJECT_ROOT, config)
    units_data, data_dictionary = _load_model_for_unit_headers()
    global_variables_data = _load_model_json("globalVariables")
    functions_data = _load_model_json("functions")
    base_path = _load_base_path()

    doc = Document()
    doc.add_heading("Software Detailed Design", 0)

    # Group by module; use data as-is from view output
    by_module = {}
    for unit_key in data.keys():
        if unit_key in ("basePath", "projectName", "unitNames"):
            continue
        unit_data = data[unit_key]
        if not isinstance(unit_data, dict) or "entries" not in unit_data:
            continue
        parts = unit_key.split(KEY_SEP, 1)
        module_name = parts[0]
        unit_name_display = unit_data.get("name", unit_key.split(KEY_SEP)[-1] if KEY_SEP in unit_key else unit_key)
        interfaces = unit_data["entries"]
        by_module.setdefault(module_name, []).append((unit_key, unit_name_display, interfaces))

    sorted_modules = sorted(by_module.keys())

    # 1 Introduction
    doc.add_heading("1 Introduction", level=1)
    doc.add_heading("1.1 Purpose", level=2)
    _add_para(doc, "[Purpose of this document.]")
    doc.add_heading("1.2 Scope", level=2)
    _add_para(doc, "[Scope of the software detailed design.]")
    doc.add_heading("1.3 Terms, Abbreviations and Definitions", level=2)
    _add_para(doc, "[Terms, abbreviations and definitions.]")

    # 2, 3, ... Modules
    n_modules = len(sorted_modules)
    for sec_idx, module_name in enumerate(sorted_modules, start=0):
        sec_num = sec_idx + 2
        print(f"  docx_exporter: {sec_idx + 1}/{n_modules} modules...", end="\r", flush=True)
        doc.add_heading(f"{sec_num} {module_name}", level=1)

        # 2.1 Static Design
        doc.add_heading(f"{sec_num}.1 Static Design", level=2)

        unit_rows_module = sorted(by_module[module_name])
        if msd_enabled and unit_rows_module:
            mod_structure_mmd = _build_module_static_structure_mermaid(module_name, unit_rows_module)
            static_mod_png = os.path.join(
                artifacts_dir, "module_static_diagrams", f"{safe_filename(module_name)}.png"
            )
            if msd_render_png:
                if _render_mermaid_to_png(PROJECT_ROOT, mod_structure_mmd, static_mod_png) and os.path.isfile(
                    static_mod_png
                ):
                    try:
                        doc.add_picture(static_mod_png, width=Inches(msd_width_in))
                    except Exception:
                        _add_mermaid_as_text(doc, mod_structure_mmd, font_small)
                else:
                    _add_mermaid_as_text(doc, mod_structure_mmd, font_small)
            else:
                _add_mermaid_as_text(doc, mod_structure_mmd, font_small)

        # Module-level index table (Component/Unit/Description/Note)
        _add_component_unit_table(
            doc,
            module_name,
            unit_rows_module,
            font_small,
            config=config,
            abbreviations=abbreviations,
        )

        unit_diag_dir = os.path.join(artifacts_dir, "unit_diagrams")
        for unit_idx, (unit_key, unit_name_display, interfaces) in enumerate(unit_rows_module, start=1):
            # 2.1.1 unit1
            doc.add_heading(f"{sec_num}.1.{unit_idx} {unit_name_display}", level=3)

            # Unit diagram (before unit header)
            unit_png = os.path.join(unit_diag_dir, f"{safe_filename(unit_key)}.png")
            if os.path.isfile(unit_png):
                try:
                    doc.add_picture(unit_png, width=Inches(6))
                except Exception:
                    _add_para(doc, f"[Unit diagram: {unit_png}]")

            # 2.1.1.1 unit header
            path_str = interfaces[0].get("location", {}).get("file", "-") if interfaces else "-"
            doc.add_heading(f"{sec_num}.1.{unit_idx}.1 unit header", level=4)
            _add_para(doc, f"Unit: {unit_name_display}")
            _add_para(doc, f"Path: {path_str}")
            unit_info = units_data.get(unit_key, {})
            unit_header_rows = _build_unit_header_table(
                unit_info,
                interfaces,
                data_dictionary,
                global_variables_data,
                base_path,
                config,
                abbreviations,
            )
            _add_unit_header_table(doc, unit_header_rows, font_small)

            # 2.1.1.2 unit interface (table)
            doc.add_heading(f"{sec_num}.1.{unit_idx}.2 unit interface", level=4)
            _add_interface_table(doc, interfaces, font_small)

            # 2.1.1.3, 2.1.1.4, ... per interface
            unit_name_flowchart = unit_key.split(KEY_SEP)[-1] if KEY_SEP in unit_key else unit_name_display
            rendered_private_fids = set()  # track private flowcharts already shown in this unit
            for iface_idx, iface in enumerate(interfaces, start=3):
                iface_id = iface.get("interfaceId", "")
                doc.add_heading(f"{sec_num}.1.{unit_idx}.{iface_idx} {unit_name_display}-{iface_id}", level=4)
                func_name = iface.get("name", "")
                unit_prefix = unit_key.replace(KEY_SEP, "_")
                flowchart = (
                    flowcharts_map.get(unit_prefix, {}).get(func_name)
                    or flowcharts_map.get(unit_name_flowchart, {}).get(func_name)
                ) if flowcharts_enabled and func_name else None
                # Build flowcharts list: own flowchart + private callee flowcharts
                flowcharts_list = []
                if flowchart:
                    png_path = None
                    if flowcharts_render_png:
                        png_path = os.path.join(flowcharts_dir, f"{unit_prefix}_{safe_filename(func_name)}.png")
                        if not os.path.isfile(png_path):
                            png_path = os.path.join(flowcharts_dir, f"{unit_name_flowchart}_{safe_filename(func_name)}.png")
                        if not os.path.isfile(png_path):
                            png_path = None
                    iface_params = ", ".join(
                        f"{p.get('type', '')} {p.get('name', '')}".strip()
                        for p in (iface.get("parameters") or [])
                    )
                    iface_return = iface.get("returnType", "") or ""
                    iface_signature = f"{iface_return} {func_name}({iface_params})".strip()
                    flowcharts_list.append((png_path, flowchart, iface_signature))

                if flowcharts_enabled:
                    callee_fids = (functions_data.get(iface.get("functionId")) or {}).get("callsIds") or []
                    for callee_fid in callee_fids:
                        callee = functions_data.get(callee_fid) or {}
                        if (callee.get("visibility") or "").lower() != "private":
                            continue
                        callee_parts = callee_fid.split(KEY_SEP)
                        callee_unit_key = KEY_SEP.join(callee_parts[:2]) if len(callee_parts) >= 2 else ""
                        callee_unit_prefix = callee_unit_key.replace(KEY_SEP, "_")
                        callee_unit_name = callee_parts[1] if len(callee_parts) > 1 else ""
                        callee_qn = callee.get("qualifiedName", "")
                        callee_func_name = callee_qn.split("::")[-1] if callee_qn else ""
                        if not callee_func_name:
                            continue
                        callee_flowchart = (
                            flowcharts_map.get(callee_unit_prefix, {}).get(callee_func_name)
                            or flowcharts_map.get(callee_unit_name, {}).get(callee_func_name)
                        )
                        if not callee_flowchart:
                            continue
                        if callee_fid in rendered_private_fids:
                            continue
                        rendered_private_fids.add(callee_fid)
                        callee_png = None
                        if flowcharts_render_png:
                            callee_png = os.path.join(flowcharts_dir, f"{callee_unit_prefix}_{safe_filename(callee_func_name)}.png")
                            if not os.path.isfile(callee_png):
                                callee_png = os.path.join(flowcharts_dir, f"{callee_unit_name}_{safe_filename(callee_func_name)}.png")
                            if not os.path.isfile(callee_png):
                                callee_png = None
                        callee_params = ", ".join(
                            f"{p.get('type', '')} {p.get('name', '')}".strip()
                            for p in (callee.get("params") or callee.get("parameters") or [])
                        )
                        callee_return = callee.get("returnType", "")
                        callee_signature = f"{callee_return} {callee_func_name}({callee_params})".strip()
                        flowcharts_list.append((callee_png, callee_flowchart, callee_signature))

                if flowcharts_list:
                    input_label = (functions_data.get(iface.get("functionId")) or {}).get("behaviourInputName") or \
                        (_readable_label(func_name) + " input").strip() if func_name else ""
                    output_label = (functions_data.get(iface.get("functionId")) or {}).get("behaviourOutputName") or \
                        (_readable_label(func_name) + " result").strip() if func_name else ""
                    _add_flowchart_table(doc, func_name, iface.get("description", ""),
                        input_label, output_label, flowcharts_list, font_small)
                else:
                    _add_para(doc, iface.get("description", "") or "-")

        # 2.2 Dynamic Behaviour: one sub-header per external call (from view output)
        doc.add_heading(f"{sec_num}.2 Dynamic Behaviour", level=2)
        docx_rows = {}
        pngs_path = os.path.join(artifacts_dir, "behaviour_diagrams", "_behaviour_pngs.json")
        if os.path.isfile(pngs_path):
            try:
                with open(pngs_path, "r", encoding="utf-8") as f:
                    docx_rows = json.load(f).get("_docxRows", {})
            except (json.JSONDecodeError, IOError):
                pass
        beh_idx = 0
        for unit_name, entries in sorted((docx_rows.get(module_name) or {}).items()):
            for row in entries:
                beh_idx += 1
                ext = row.get("externalUnitFunction", "")
                current_fn = row.get("currentFunctionName", "") or ""
                subheader = f"{unit_name} - {current_fn}"
                if ext:
                    subheader += f" ({ext})"
                doc.add_heading(f"{sec_num}.2.{beh_idx} {subheader}", level=3)
                # Prefer precomputed behaviourInputName / behaviourOutputName from model_deriver
                input_label = ""
                output_label = ""
                try:
                    for fid, f in (functions_data or {}).items():
                        parts = fid.split("|")
                        if len(parts) < 3:
                            continue
                        mod, unit, _ = parts[0], parts[1], parts[2]
                        if mod != module_name or unit != unit_name:
                            continue
                        qn = f.get("qualifiedName", "") or ""
                        base_name = qn.split("::")[-1] if qn else ""
                        if base_name != current_fn:
                            continue
                        input_label = (f.get("behaviourInputName") or "").strip()
                        output_label = (f.get("behaviourOutputName") or "").strip()
                        break
                except Exception:
                    input_label = input_label or ""
                    output_label = output_label or ""

                # Fallback if model is old/missing fields
                if not input_label:
                    base_fn_label = _readable_label(current_fn)
                    input_label = (base_fn_label + " input").strip() if base_fn_label else "Behaviour input"
                if not output_label:
                    base_fn_label = _readable_label(current_fn)
                    output_label = (base_fn_label + " result").strip() if base_fn_label else "Behaviour result"

                _add_behavior_description_table(doc, row.get("behaviorDescription", None), input_label, output_label)
                p = doc.add_paragraph()
                r = p.add_run("Behaviour")
                r.bold = True
                png_path = row.get("pngPath")
                if png_path and os.path.isfile(png_path):
                    try:
                        doc.add_picture(png_path, width=Inches(6))
                    except Exception:
                        _add_para(doc, f"[Behaviour diagram: {png_path}]")
                elif png_path:
                    _add_para(doc, f"[Behaviour diagram: {png_path}]")

    print()  # newline after progress
    # N Code Metrics, Coding rule, test coverage
    metrics_sec = len(sorted_modules) + 2
    doc.add_heading(f"{metrics_sec} Code Metrics, Coding Rule, Test Coverage", level=1)
    _add_para(doc, "[Code metrics, coding rules and test coverage.]")

    # Appendix A
    doc.add_heading("Appendix A. Design Guideline", level=1)
    _add_para(doc, "[Design guidelines.]")

    os.makedirs(os.path.dirname(docx_path) or ".", exist_ok=True)
    doc.save(docx_path)
    return (True, docx_path)


def main():
    args = sys.argv[1:]
    selected_group = None
    if "--selected-group" in args:
        i = args.index("--selected-group")
        if i + 1 < len(args):
            selected_group = args[i + 1]
            # Remove the flag + value so positional args remain: json_path docx_path
            args = args[:i] + args[i + 2 :]

    json_path = args[0] if len(args) >= 1 else None
    docx_path = args[1] if len(args) >= 2 else None
    ok, out_path = export_docx(json_path, docx_path, selected_group=selected_group)
    if ok:
        print(f"Exported: {out_path}")
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()

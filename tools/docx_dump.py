#!/usr/bin/env python3
"""Dump a generated .docx to deterministic plain text, for review and diffing.

The DOCX exporter interleaves headings, paragraphs, tables, and diagram images.
`python-docx` exposes `doc.paragraphs` and `doc.tables` as separate flat lists,
which loses that interleaving — so this walks the body XML directly and keeps
everything in document order.

    python tools/docx_dump.py <docx>
    python tools/docx_dump.py <docx> -o dump.md

Diagrams are emitted as `[image w=6.00in sha=1a2b3c4d]`: the sha is a hash of
the image bytes, so an unchanged diagram gives an identical line while a
regenerated one shows up as a change. Nothing in the output depends on run
order or file paths, so two dumps are directly comparable:

    python tools/docx_dump.py poc4.docx > /tmp/poc4.md
    python tools/docx_dump.py db.docx   > /tmp/db.md
    diff -u /tmp/poc4.md /tmp/db.md
"""
from __future__ import annotations

import argparse
import hashlib
import os
import re
import sys

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:  # pragma: no cover - dependency is in requirements.txt
    sys.exit("python-docx is not installed: pip install python-docx")

EMU_PER_INCH = 914400

_HEADING_RE = re.compile(r"^Heading (\d+)$")

# Multi-paragraph table cells are flattened onto one line so the row stays a
# single markdown row; the marker keeps the internal breaks visible.
CELL_BREAK = " // "


def _heading_level(para: Paragraph) -> int | None:
    """Heading depth for a paragraph, or None if it is body text."""
    try:
        name = para.style.name or ""
    except Exception:
        return None
    m = _HEADING_RE.match(name)
    return int(m.group(1)) if m else None


def _images(para: Paragraph) -> list[str]:
    """Render every drawing in a paragraph as a stable `[image ...]` marker."""
    out = []
    for drawing in para._p.iter(qn("w:drawing")):
        width = ""
        extent = drawing.find(".//" + qn("wp:extent"))
        if extent is not None and extent.get("cx"):
            width = " w=%.2fin" % (int(extent.get("cx")) / EMU_PER_INCH)
        sha = ""
        blip = drawing.find(".//" + qn("a:blip"))
        if blip is not None:
            rid = blip.get(qn("r:embed"))
            try:
                blob = para.part.related_parts[rid].blob
                sha = " sha=" + hashlib.sha1(blob).hexdigest()[:8]
            except Exception:
                sha = " sha=?"
        out.append(f"[image{width}{sha}]")
    return out


def _para_lines(para: Paragraph) -> list[str]:
    """One paragraph as zero or more output lines."""
    lines = []
    text = " ".join(para.text.split())
    level = _heading_level(para)
    if level:
        lines.append("")
        lines.append("#" * min(level, 6) + " " + text)
    elif text:
        lines.append(text)
    lines.extend(_images(para))
    if para._p.find(".//" + qn("w:br") + '[@{%s}type="page"]' % para._p.nsmap["w"]) is not None:
        lines.append("[page break]")
    return lines


def _cell_text(cell) -> str:
    """Flatten a cell — paragraphs, images, and any nested table — to one line."""
    parts = []
    for child in cell._tc.iterchildren():
        if child.tag == qn("w:p"):
            para = Paragraph(child, cell)
            text = " ".join(para.text.split())
            if text:
                parts.append(text)
            parts.extend(_images(para))
        elif child.tag == qn("w:tbl"):
            table = Table(child, cell)
            parts.append(f"[nested table {len(table.rows)}x{len(table.columns)}]")
    return CELL_BREAK.join(parts).replace("|", r"\|")


def _table_lines(table: Table) -> list[str]:
    """A table as markdown pipe rows, with a separator after the header row."""
    lines = [""]
    for i, row in enumerate(table.rows):
        cells = [_cell_text(c) for c in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
        if i == 0:
            lines.append("|" + "|".join(["---"] * len(cells)) + "|")
    return lines


def dump(path: str) -> str:
    doc = Document(path)
    # No filename header: the two sides of a diff are usually named differently
    # (poc4.docx vs db.docx) and a header would show as a change on every comparison.
    lines: list[str] = []
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            lines.extend(_para_lines(Paragraph(child, doc)))
        elif child.tag == qn("w:tbl"):
            lines.extend(_table_lines(Table(child, doc)))
    # Collapse the runs of blank lines the heading/table spacing introduces.
    out, blank = [], False
    for line in lines:
        if not line.strip():
            if blank:
                continue
            blank = True
        else:
            blank = False
        out.append(line.rstrip())
    return "\n".join(out).strip() + "\n"


def main() -> int:
    ap = argparse.ArgumentParser(description="Dump a .docx to plain text for review and diffing.")
    ap.add_argument("docx", help="path to the .docx file")
    ap.add_argument("-o", "--out", help="write to this file instead of stdout")
    args = ap.parse_args()

    if not os.path.isfile(args.docx):
        sys.exit(f"not a file: {args.docx}")

    text = dump(args.docx)
    if args.out:
        with open(args.out, "w", encoding="utf-8") as fh:
            fh.write(text)
        print(f"wrote {args.out} ({len(text.splitlines())} lines)")
    else:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stdout.write(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

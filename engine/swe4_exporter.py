"""Export test_specs.json -> Software Unit Test Specification DOCX (SWE.4).

Phase-4 exporter for the swe4 doc type (dispatched via EXPORTER_REGISTRY). Reads
the testSpecs view output and renders, per function, the two tables of
docs/spec/SWE4_WIKI.md:

  Table A (horizontal, 1 header + 1 data row) -- the test *content*:
    Eval. Equipment Name | Precondition | Input | Test Steps | Expected Results |
    Test Platform
  Table B (vertical form) -- the *metadata*:
    Test Case ID | Alias Test ID | Priority | Risk | Test Method |
    Test Environment | Test Case Generation Method | Linked Work Items

Precondition, Input and Expected Results are flat numbered lists; only Test Steps
nests, because it follows the control flow. Cover/TOC/intro helpers are shared
with the SWE.3 exporter via docx_common.
"""
import os
import sys
import json
from typing import Optional, Tuple

from utils import KEY_SEP
from core.paths import paths as _paths
from docx_common import (
    load_abbreviations, set_cell_font, add_para, add_toc, build_cover_page,
)

_p = _paths()
PROJECT_ROOT = _p.project_root
OUTPUT_DIR = _p.output_dir
MODEL_DIR = _p.model_dir

TABLE_A_COLS = ("Eval. Equipment Name", "Precondition", "Input",
                "Test Steps", "Expected Results", "Test Platform")

_INDENT = "    "


# ---------------------------------------------------------------------------
# Cell renderers
# ---------------------------------------------------------------------------

def _numbered(items):
    """Flat `1) ... 2) ...` list -- Precondition, Input, Expected Results."""
    return "\n".join(f"{i}) {t}" for i, t in enumerate(items, start=1))


def _precondition_text(pre: dict) -> str:
    """Three entries: mock functions, parameters, globals -- each a comma-separated
    list, names and types only (no values, no ranges)."""
    entries = []
    mocks = pre.get("mockFunctions") or []
    if mocks:
        entries.append("Mock functions: " + ", ".join(mocks))
    params = pre.get("parameters") or []
    if params:
        entries.append("Parameters: " + ", ".join(p.get("text", "") for p in params))
    globals_ = pre.get("globals") or []
    if globals_:
        entries.append("Globals: " + ", ".join(g.get("text", "") for g in globals_))
    return _numbered(entries) if entries else "None"


def _input_text(inp: dict) -> str:
    """One entry per input, `type variable[low-high]`. VOID when nothing is read."""
    entries = inp.get("entries") or []
    if not entries:
        return "VOID"
    return _numbered([e.get("text", "") for e in entries])


def _expected_text(exp: dict) -> str:
    """Assertions: mocks first, then one entry per return (naming its step),
    then out-parameters and written globals."""
    entries = []
    mocks = exp.get("mockFunctions") or []
    if mocks:
        entries.append("Successfully called mock functions " + ", ".join(mocks))
    for r in exp.get("returns") or []:
        step = r.get("step")
        entries.append(f"{r.get('text', '')}" + (f" in step {step}" if step else ""))
    for o in exp.get("outParameters") or []:
        entries.append(f"Successfully updated {o.get('text', '')}")
    for g in exp.get("globals") or []:
        entries.append(f"Successfully updated {g.get('text', '')}")
    if not entries:
        return "No return value; no global side effects"
    return _numbered(entries)


def _test_steps_text(spec: dict) -> str:
    """Nested numbered transcription; depth is shown by indentation."""
    steps = spec.get("testSteps") or []
    if not steps:
        return "Not available (no control-flow graph for this function)"
    lines = []
    for st in steps:
        number = st.get("number", "")
        depth = number.count(".")
        lines.append(f"{_INDENT * depth}{number}) {st.get('text', '')}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

def _add_table_a(doc, spec: dict, cfg_swe4: dict, font_small) -> None:
    table = doc.add_table(rows=1, cols=len(TABLE_A_COLS))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(TABLE_A_COLS):
        hdr[i].text = c
        set_cell_font(hdr[i], font_small, bold=True)
    row = table.add_row().cells
    values = (
        cfg_swe4.get("evalEquipmentName", "Emulator"),
        _precondition_text(spec.get("precondition", {})),
        _input_text(spec.get("input", {})),
        _test_steps_text(spec),
        _expected_text(spec.get("expected", {})),
        cfg_swe4.get("testPlatform", "VectorCAST"),
    )
    for i, txt in enumerate(values):
        row[i].text = str(txt)
        set_cell_font(row[i], font_small)


def _add_table_b(doc, spec: dict, cfg_swe4: dict, font_small) -> None:
    fields = [
        ("Test Case ID", spec.get("testCaseId", "")),
        ("Alias Test ID", "-"),
        ("Priority", cfg_swe4.get("priorityDefault", "Medium")),
        ("Risk", "-"),
        ("Test Method", "-"),
        ("Test Environment", cfg_swe4.get("testEnvironment", "Emulator")),
        ("Test Case Generation Method", spec.get("generationMethod", "")),
        ("Linked Work Items", "-"),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    for r, (label, value) in enumerate(fields):
        cells = table.rows[r].cells
        cells[0].text = label
        cells[1].text = str(value)
        set_cell_font(cells[0], font_small, bold=True)
        set_cell_font(cells[1], font_small)


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

def export_test_specs(json_path: str = None, docx_path: str = None,
                      selected_group: str = None,
                      selected_components: list = None) -> Tuple[bool, Optional[str]]:
    from core.config import app_config
    config = app_config()
    json_path = os.path.abspath(json_path or os.path.join(OUTPUT_DIR, "test_specs.json"))

    if selected_components:
        group_name = "_".join(selected_components)
    else:
        group_name = selected_group or "all"
    if not docx_path:
        docx_path = os.path.join(
            PROJECT_ROOT, f"output/software_unit_test_specification_{group_name}.docx")

    if not os.path.isfile(json_path):
        print(f"Error: {json_path} not found. Run the testSpecs view (Phase 3) first.")
        return (False, None)

    try:
        from docx import Document
        from docx.shared import Pt
        font_small = Pt(8)
    except ImportError:
        print("Error: python-docx not installed. pip install python-docx")
        return (False, None)

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    docx_cfg = config.get("docx", {}) or {}
    cfg_swe4 = docx_cfg.get("swe4", {}) or {}
    intro_cfg = cfg_swe4.get("introduction", {}) or {}
    abbreviations = load_abbreviations(PROJECT_ROOT, config)

    project_name = "Software Project"
    _meta_path = os.path.join(MODEL_DIR, "metadata.json")
    if os.path.isfile(_meta_path):
        with open(_meta_path, "r", encoding="utf-8") as _f:
            project_name = json.load(_f).get("projectName", project_name)

    from core.config import get_group_layer_name, get_component_layer_name
    if selected_components:
        _layer = get_component_layer_name(config, selected_components[0])
        _disp = " / ".join(c.replace("-", " ") for c in selected_components)
        cover_group = f"{_layer} {_disp}" if _layer else _disp
    elif selected_group:
        _layer = get_group_layer_name(config, selected_group)
        _disp = selected_group.replace("-", " ")
        cover_group = f"{_layer} {_disp}" if _layer else _disp
    else:
        cover_group = "All Units"

    # Group the per-unit specs by component (unit_key = "Component|Unit").
    by_component = {}
    for unit_key, unit_data in data.items():
        if unit_key == "unitNames" or not isinstance(unit_data, dict):
            continue
        functions = unit_data.get("functions") or []
        if not functions:
            continue
        component_name = unit_key.split(KEY_SEP, 1)[0]
        unit_name = unit_data.get("name", unit_key.split(KEY_SEP)[-1])
        by_component.setdefault(component_name, []).append((unit_key, unit_name, functions))
    sorted_components = sorted(by_component.keys())

    doc = Document()
    build_cover_page(doc, project_name, cover_group,
                     copyright_text=docx_cfg.get("copyrightText", ""),
                     subtitle_prefix="Software Unit Test Specification")
    add_toc(doc)

    # 1 Introduction
    doc.add_heading("1 Introduction", level=1)
    doc.add_heading("1.1 Purpose", level=2)
    add_para(doc, intro_cfg.get("purpose", "[Purpose of this document.]")
             .replace("{project_name}", project_name))
    doc.add_heading("1.2 Scope", level=2)
    add_para(doc, intro_cfg.get("scopeIntro", "[Scope of the unit test specification.]")
             .replace("{project_name}", project_name))
    for _comp in sorted_components:
        add_para(doc, f"• {_comp.replace('-', ' ')}")
    _scope_body = intro_cfg.get("scopeBody", "")
    if _scope_body:
        add_para(doc, _scope_body)
    doc.add_heading("1.3 Terms, Abbreviations and Definitions", level=2)
    if abbreviations:
        _tbl = doc.add_table(rows=1, cols=2)
        _tbl.style = "Table Grid"
        _hdr = _tbl.rows[0].cells
        _hdr[0].text = "Term"
        _hdr[1].text = "Description"
        for _cell in _hdr:
            for _run in _cell.paragraphs[0].runs:
                _run.bold = True
        for _term, _desc in sorted(abbreviations.items()):
            _r = _tbl.add_row().cells
            _r[0].text = _term
            _r[1].text = _desc
    else:
        add_para(doc, "[Terms, abbreviations and definitions.]")

    # 2 Software Unit Test Specification
    doc.add_heading("2 Software Unit Test Specification", level=1)
    spec_total = 0
    for comp_idx, component_name in enumerate(sorted_components, start=1):
        sec = f"2.{comp_idx}"
        doc.add_heading(f"{sec} {component_name.replace('-', ' ')}", level=2)
        units = sorted(by_component[component_name], key=lambda t: t[1])
        for unit_idx, (unit_key, unit_name, functions) in enumerate(units, start=1):
            doc.add_heading(f"{sec}.{unit_idx} {unit_name}", level=3)
            for fn_idx, spec in enumerate(
                    sorted(functions, key=lambda s: s.get("location", {}).get("line", 0)),
                    start=1):
                fn_name = spec.get("name", "")
                doc.add_heading(f"{sec}.{unit_idx}.{fn_idx} {unit_name}-{fn_name}", level=4)
                if spec.get("description"):
                    add_para(doc, spec["description"])
                _add_table_a(doc, spec, cfg_swe4, font_small)
                add_para(doc, "")
                _add_table_b(doc, spec, cfg_swe4, font_small)
                spec_total += 1

    # 3 Code Metric, Coding Rule, Test Coverage (placeholder -- coverage tooling parked)
    doc.add_heading("3 Code Metric, Coding Rule, Test Coverage", level=1)
    add_para(doc, "[Code metrics, coding-rule compliance, and test-coverage results.]")

    doc.add_heading("Appendix A. Reference", level=1)

    os.makedirs(os.path.dirname(os.path.abspath(docx_path)), exist_ok=True)
    doc.save(docx_path)
    from utils import log
    log("%s (%d components, %d function specs)" % (docx_path, len(sorted_components), spec_total),
        component="swe4_exporter")
    return (True, docx_path)


def main():
    args = sys.argv[1:]
    selected_group = None
    if "--selected-group" in args:
        i = args.index("--selected-group")
        if i + 1 < len(args):
            selected_group = args[i + 1]
            args = args[:i] + args[i + 2:]

    selected_components = []
    filtered_args = []
    j = 0
    while j < len(args):
        if args[j] == "--selected-component" and j + 1 < len(args):
            selected_components.append(args[j + 1])
            j += 2
        else:
            filtered_args.append(args[j])
            j += 1
    args = filtered_args

    json_path = args[0] if len(args) >= 1 else None
    docx_path = args[1] if len(args) >= 2 else None
    ok, out_path = export_test_specs(json_path, docx_path,
                                     selected_group=selected_group,
                                     selected_components=selected_components or None)
    if ok:
        print(f"Exported: {out_path}")
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()

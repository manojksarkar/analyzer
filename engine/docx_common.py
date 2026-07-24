"""Shared DOCX building blocks for the document exporters.

Leaf helpers used by more than one exporter (docx_exporter.py for SWE.3,
swe4_exporter.py for SWE.4, and api/services/doc_render.py). Keeping them here
stops the exporters from drifting. Nothing here knows about a specific document
type — callers pass the doc + data; the doc-type-specific structure lives in
each exporter.
"""
import os
import json

from core.paths import paths as _paths

_p = _paths()
PROJECT_ROOT = _p.project_root
MODEL_DIR = _p.model_dir


# --------------------------------------------------------------------------
# Model / config loaders
# --------------------------------------------------------------------------
def load_model_json(name: str) -> dict:
    path = os.path.join(MODEL_DIR, f"{name}.json")
    if not os.path.isfile(path):
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        return {}


def load_base_path() -> str:
    meta = load_model_json("metadata")
    return (meta.get("basePath") or "").strip()


def load_abbreviations(project_root: str, config: dict) -> dict:
    """Load abbreviations from the file named in config (llm.abbreviationsPath).

    Format: one per line, 'abbrev: meaning' or 'abbrev=meaning'; # = comment.
    """
    path = (config.get("llm") or {}).get("abbreviationsPath", "").strip()
    if not path:
        return {}
    full_path = os.path.join(project_root, path) if not os.path.isabs(path) else path
    if not os.path.isfile(full_path):
        return {}
    result = {}
    try:
        with open(full_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#"):
                    continue
                if ":" in line:
                    k, _, v = line.partition(":")
                elif "=" in line:
                    k, _, v = line.partition("=")
                else:
                    continue
                k, v = k.strip(), v.strip()
                if k:
                    result[k] = v
        return result
    except OSError:
        return {}


# --------------------------------------------------------------------------
# Low-level docx helpers
# --------------------------------------------------------------------------
def set_cell_font(cell, font_pt, bold=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = font_pt
            r.font.bold = bold


def add_horizontal_rule(doc) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_para(doc, text, style="Normal"):
    return doc.add_paragraph(text, style=style)


def add_toc(doc) -> None:
    """Insert a Word automatic table of contents field followed by a page break."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    try:
        p_title = doc.add_paragraph("Contents", style="TOC Heading")
    except KeyError:
        from docx.shared import Pt
        p_title = doc.add_paragraph()
        run = p_title.add_run("Contents")
        run.font.size = Pt(16)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_title.runs:
        run.bold = True

    p = doc.add_paragraph()

    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    fldChar.set(qn("w:dirty"), "true")
    run._r.append(fldChar)

    run2 = p.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-4" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run3._r.append(fldChar2)

    run4 = p.add_run()
    t = OxmlElement("w:t")
    t.text = "Right-click here and select 'Update Field' to populate the table of contents."
    run4._r.append(t)

    run5 = p.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run5._r.append(fldChar3)

    doc.add_page_break()

    # Tell Word to update all fields (including this TOC) when the document is opened
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings.element.append(update_fields)


def build_cover_page(doc, project_name: str, group_name: str, version: str = "1.0.0",
                     copyright_text: str = "",
                     subtitle_prefix: str = "Software Detailed Design Specification") -> None:
    """Render the cover page (first page) of the DOCX.

    `subtitle_prefix` names the document kind on the cover; it defaults to the
    SWE.3 title so existing callers are unchanged, and SWE.4 passes its own.
    """
    from datetime import date as _date
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY = RGBColor(30, 60, 120)
    DARK = RGBColor(60, 60, 60)
    ASSETS = os.path.join(PROJECT_ROOT, "engine", "assets")

    def _spacing(para, before=0, after=0):
        pPr = para._p.get_or_add_pPr()
        for old in pPr.findall(qn("w:spacing")):
            pPr.remove(old)
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), str(before))
        sp.set(qn("w:after"),  str(after))
        pPr.append(sp)

    def _align(para, val="right"):
        pPr = para._p.get_or_add_pPr()
        for old in pPr.findall(qn("w:jc")):
            pPr.remove(old)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), val)
        pPr.append(jc)
        para.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if val == "right" else
                          WD_ALIGN_PARAGRAPH.CENTER if val == "center" else
                          WD_ALIGN_PARAGRAPH.LEFT)

    def _run(para, text, size_pt, bold=False, color=None):
        r = para.add_run(text)
        r.bold = bold
        r.font.size = Pt(size_pt)
        if color:
            r.font.color.rgb = color
        return r

    def _double_underline(run, color_hex="1E3C78"):
        rPr = run._r.get_or_add_rPr()
        for old in rPr.findall(qn("w:u")):
            rPr.remove(old)
        u = OxmlElement("w:u")
        u.set(qn("w:val"),   "double")
        u.set(qn("w:color"), color_hex)
        u.set(qn("w:sz"),    "12")
        rPr.append(u)

    def _spacer(n=1):
        for _ in range(n):
            p = doc.add_paragraph()
            _spacing(p, 0, 0)

    section    = doc.sections[0]
    body_w_in  = (section.page_width / 914400) - (
        section.left_margin / 914400 + section.right_margin / 914400)

    _spacer(8)

    # Project name — largest, bold, double-underlined
    p_name = doc.add_paragraph()
    _spacing(p_name, before=0, after=120)
    _align(p_name, "right")
    r_name = _run(p_name, project_name, size_pt=36, bold=True, color=NAVY)
    _double_underline(r_name, "1E3C78")

    # Subtitle — single line, no dash
    p_sub = doc.add_paragraph()
    _spacing(p_sub, before=0, after=100)
    _align(p_sub, "right")
    _run(p_sub, f"{subtitle_prefix}  {group_name}", size_pt=16, bold=True, color=NAVY)

    # Version
    p_ver = doc.add_paragraph()
    _spacing(p_ver, before=0, after=60)
    _align(p_ver, "right")
    _run(p_ver, f"Version {version}", size_pt=12, color=DARK)

    # Date
    p_date = doc.add_paragraph()
    _spacing(p_date, before=0, after=400)
    _align(p_date, "right")
    _run(p_date, _date.today().strftime("%Y-%m-%d"), size_pt=12, color=DARK)

    # Copyright image — left-aligned
    cr_path = os.path.join(ASSETS, "copyright.png")
    p_cr = doc.add_paragraph()
    _spacing(p_cr, before=0, after=0)
    _align(p_cr, "left")
    if os.path.isfile(cr_path):
        p_cr.add_run().add_picture(cr_path, width=Inches(2.6))
    else:
        _run(p_cr, "© All Rights Reserved", size_pt=10, color=DARK)

    # Copyright sentence below the image
    _cr_text = copyright_text or f"© {_date.today().year} All Rights Reserved."
    p_cr_text = doc.add_paragraph()
    _spacing(p_cr_text, before=0, after=0)
    _align(p_cr_text, "left")
    _run(p_cr_text, _cr_text, size_pt=8, color=RGBColor(128, 128, 128))

    _spacer(4)

    # Bottom arc — full body width
    arc_path = os.path.join(ASSETS, "bottom_arc.png")
    p_arc = doc.add_paragraph()
    _spacing(p_arc, before=0, after=0)
    _align(p_arc, "center")
    if os.path.isfile(arc_path):
        p_arc.add_run().add_picture(arc_path, width=Inches(body_w_in))

    doc.add_page_break()

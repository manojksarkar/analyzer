"""SWE.4 exporter: cell rendering and the view -> DOCX chain.

The cell renderers are the contract of docs/spec/SWE4_WIKI.md: Precondition,
Input and Expected Results are flat `1) 2) 3)` lists; only Test Steps nests.
"""
import os
import sys

import pytest

ENGINE = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(
    os.path.abspath(__file__)))), "engine")
if ENGINE not in sys.path:
    sys.path.insert(0, ENGINE)

from swe4_exporter import (  # noqa: E402
    _precondition_text, _input_text, _expected_text, _test_steps_text,
    TABLE_A_COLS, export_test_specs,
)


# --- Precondition ----------------------------------------------------------

def test_precondition_is_three_numbered_entries():
    text = _precondition_text({
        "mockFunctions": ["a()", "b()"],
        "parameters": [{"text": "int x"}],
        "globals": [{"text": "char g"}],
    })
    assert text.splitlines() == [
        "1) Mock functions: a(), b()",
        "2) Parameters: int x",
        "3) Globals: char g",
    ]


def test_precondition_omits_empty_groups_and_renumbers():
    text = _precondition_text({"mockFunctions": [], "parameters": [{"text": "int x"}],
                               "globals": []})
    assert text == "1) Parameters: int x"


def test_precondition_with_nothing_reads_none():
    assert _precondition_text({}) == "None"


# --- Input -----------------------------------------------------------------

def test_input_is_a_flat_numbered_list():
    text = _input_text({"entries": [{"text": "int a[0-9]"}, {"text": "char b[0-1]"}]})
    assert text.splitlines() == ["1) int a[0-9]", "2) char b[0-1]"]


def test_input_is_void_when_there_are_no_entries():
    assert _input_text({"entries": []}) == "VOID"


def test_input_never_nests():
    text = _input_text({"entries": [{"text": "int a[0-9]"}, {"text": "int b[0-9]"}]})
    assert not any(line.startswith(("  ", "1.1")) for line in text.splitlines())


# --- Expected Results ------------------------------------------------------

def test_expected_puts_mocks_first_then_returns_with_step_refs():
    text = _expected_text({
        "mockFunctions": ["m()"],
        "returns": [{"text": "Successfully returned -1", "step": "2.2"},
                    {"text": "Successfully returned 0", "step": "3"}],
        "outParameters": [{"text": "int* out[0-9]"}],
        "globals": [{"text": "char g[0-1]"}],
    })
    assert text.splitlines() == [
        "1) Successfully called mock functions m()",
        "2) Successfully returned -1 in step 2.2",
        "3) Successfully returned 0 in step 3",
        "4) Successfully updated int* out[0-9]",
        "5) Successfully updated char g[0-1]",
    ]


def test_expected_states_when_nothing_changes():
    assert _expected_text({}) == "No return value; no global side effects"


# --- Test Steps ------------------------------------------------------------

def test_test_steps_indent_by_depth():
    text = _test_steps_text({"testSteps": [
        {"number": "1", "text": "Issue function fn."},
        {"number": "2", "text": "Check whether x."},
        {"number": "2.1", "text": "True: Return 1."},
        {"number": "2.1.1", "text": "Deeper."},
    ]})
    lines = text.splitlines()
    assert lines[0] == "1) Issue function fn."
    assert lines[2].startswith("    2.1)")
    assert lines[3].startswith("        2.1.1)")


def test_missing_steps_say_so_rather_than_rendering_empty():
    assert "no control-flow graph" in _test_steps_text({"testSteps": []})


# --- Document chain --------------------------------------------------------

def test_table_a_has_the_six_wiki_columns():
    assert TABLE_A_COLS == ("Eval. Equipment Name", "Precondition", "Input",
                            "Test Steps", "Expected Results", "Test Platform")


@pytest.fixture
def spec_json(tmp_path):
    import json
    data = {
        "unitNames": {"C|U": "U"},
        "C|U": {"name": "U", "functions": [{
            "functionId": "C|U|fn", "interfaceId": "IF_01", "testCaseId": "TC_IF_01",
            "name": "fn", "qualifiedName": "fn", "unitKey": "C|U", "unitName": "U",
            "location": {"file": "U.cpp", "line": 1}, "returnType": "int",
            "generationMethod": "Analysis of Requirements",
            "precondition": {"mockFunctions": ["m()"],
                             "parameters": [{"text": "int x"}], "globals": []},
            "input": {"entries": [{"text": "int x[0-9]"}], "isVoid": False},
            "expected": {"mockFunctions": ["m()"],
                         "returns": [{"text": "Successfully returned 0", "step": "2"}],
                         "outParameters": [], "globals": []},
            "testSteps": [{"number": "1", "text": "Issue function fn with inputs x."},
                          {"number": "2", "text": "Return 0."}],
        }]},
    }
    p = tmp_path / "test_specs.json"
    p.write_text(json.dumps(data), encoding="utf-8")
    return p


def test_view_output_exports_to_a_readable_docx(spec_json, tmp_path):
    pytest.importorskip("docx")
    from docx import Document
    out = tmp_path / "swe4.docx"
    ok, path = export_test_specs(str(spec_json), str(out), selected_group="G")
    assert ok and os.path.isfile(path)

    doc = Document(path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert headings[0] == "1 Introduction"
    assert "2 Software Unit Test Specification" in headings
    assert "3 Code Metric, Coding Rule, Test Coverage" in headings
    assert headings[-1] == "Appendix A. Reference"
    assert any(h.endswith("U-fn") for h in headings)


def test_each_function_gets_one_table_a_and_one_table_b(spec_json, tmp_path):
    pytest.importorskip("docx")
    from docx import Document
    ok, path = export_test_specs(str(spec_json), str(tmp_path / "s.docx"),
                                 selected_group="G")
    assert ok
    doc = Document(path)
    table_a = [t for t in doc.tables if t.rows[0].cells[0].text == TABLE_A_COLS[0]]
    table_b = [t for t in doc.tables if t.rows[0].cells[0].text == "Test Case ID"]
    assert len(table_a) == 1 and len(table_b) == 1
    assert len(table_a[0].rows) == 2          # one header + one data row
    assert len(table_b[0].rows) == 8          # eight metadata fields


def test_missing_input_json_fails_cleanly(tmp_path):
    ok, path = export_test_specs(str(tmp_path / "nope.json"), str(tmp_path / "o.docx"))
    assert ok is False and path is None

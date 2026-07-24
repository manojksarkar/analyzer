"""End-to-end doc chain for SWE.4: view JSON -> exporter -> .docx.

Chains the real testSpecs view builder and the real swe4 exporter through the
on-disk test_specs.json contract, then opens the generated document and asserts
its structure. Self-contained (no pipeline, model, or LLM) so it guards the
view<->exporter schema against drift in CI. Structural, not golden-binary —
mirrors the repo's e2e docx convention.
"""
import os
import sys
import json

import pytest

pytestmark = pytest.mark.unit

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, os.path.join(PROJECT_ROOT, "engine"))

pytest.importorskip("docx")

from views.test_specs import _build_test_specs
from swe4_exporter import export_test_specs, TABLE_A_COLS


def _model():
    functions = {
        "Comp|U|doWrite|int": {
            "qualifiedName": "doWrite", "interfaceId": "IF_01", "visibility": "public",
            "returnType": "void", "parameters": [{"name": "n", "type": "int"}],
            "location": {"line": 10}, "callsIds": ["Comp|U|helper|"],
            "writesGlobalIds": ["Comp|U|g_count"], "readsGlobalIds": ["Comp|U|g_count"],
        },
        "Comp|U|readOnly|": {
            "qualifiedName": "readOnly", "interfaceId": "IF_02", "visibility": "public",
            "returnType": "int", "parameters": [], "location": {"line": 20},
        },
        "Comp|U|helper|": {
            "qualifiedName": "helper", "interfaceId": "IF_03", "visibility": "private",
            "returnType": "void", "parameters": [], "location": {"line": 30},
        },
    }
    globals_ = {"Comp|U|g_count": {"qualifiedName": "g_count", "type": "int", "value": "0"}}
    units = {"Comp|U": {"name": "U", "fileName": "U.cpp",
                        "functionIds": list(functions.keys())}}
    return units, functions, globals_


@pytest.fixture(scope="module")
def rendered_docx(tmp_path_factory):
    from docx import Document
    out = tmp_path_factory.mktemp("swe4")
    specs = _build_test_specs(*_model())
    json_path = os.path.join(out, "test_specs.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(specs, f, indent=2)
    docx_path = os.path.join(out, "uts.docx")
    ok, produced = export_test_specs(json_path, docx_path)
    assert ok and os.path.isfile(produced)
    return Document(produced)


def _headings(doc):
    return [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]


class TestDocumentStructure:
    def test_required_sections_in_order(self, rendered_docx):
        heads = _headings(rendered_docx)
        for expected in ["1 Introduction",
                         "2 Software Unit Test Specification",
                         "3 Code Metric, Coding Rule, Test Coverage",
                         "Appendix A. Reference"]:
            assert expected in heads, f"missing section: {expected}"
        # order preserved
        idx = [heads.index(h) for h in
               ["1 Introduction", "2 Software Unit Test Specification",
                "3 Code Metric, Coding Rule, Test Coverage", "Appendix A. Reference"]]
        assert idx == sorted(idx)

    def test_public_function_headings_present_private_absent(self, rendered_docx):
        heads = " | ".join(_headings(rendered_docx))
        assert "doWrite" in heads and "readOnly" in heads
        assert "helper" not in heads  # private -> no spec


class TestTables:
    def test_table_a_horizontal_six_cols(self, rendered_docx):
        a_tables = [t for t in rendered_docx.tables
                    if [c.text for c in t.rows[0].cells] == list(TABLE_A_COLS)]
        assert a_tables, "no Table A found"
        for t in a_tables:
            assert len(t.columns) == 6 and len(t.rows) == 2  # header + 1 data row

    def test_table_b_vertical_metadata(self, rendered_docx):
        b_tables = [t for t in rendered_docx.tables
                    if len(t.columns) == 2 and t.rows and t.rows[0].cells[0].text == "Test Case ID"]
        assert b_tables, "no Table B found"
        for t in b_tables:
            labels = [r.cells[0].text for r in t.rows]
            assert "Test Case Generation Method" in labels
            gen = next(r for r in t.rows if r.cells[0].text == "Test Case Generation Method")
            assert gen.cells[1].text == "Analysis of Requirements"

    def test_one_table_a_and_b_per_public_function(self, rendered_docx):
        a = [t for t in rendered_docx.tables
             if [c.text for c in t.rows[0].cells] == list(TABLE_A_COLS)]
        b = [t for t in rendered_docx.tables
             if len(t.columns) == 2 and t.rows and t.rows[0].cells[0].text == "Test Case ID"]
        assert len(a) == 2 and len(b) == 2  # two public functions

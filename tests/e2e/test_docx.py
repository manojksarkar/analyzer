"""E2E tests — verifies the final output/software_detailed_design_Sample.docx.

Column index map for interface tables (COLS in docx_exporter.py):
  0  Interface ID
  1  Interface Name
  2  Information
  3  Data Type
  4  Data Range
  5  Direction(In/Out)
  6  Source/Destination
  7  Interface Type
"""
import os

import pytest

try:
    from docx import Document
except ImportError:
    pytest.skip("python-docx not installed", allow_module_level=True)

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DOCX_PATH = os.path.join(PROJECT_ROOT, "output", "software_detailed_design_Sample.docx")

COL_IF_ID    = 0
COL_IF_NAME  = 1
COL_DIRECTION = 5
COL_IF_TYPE  = 7

PRIVATE_NAMES = {"coreHelper", "coreSwitch", "libClamp", "utilClip", "g_count"}


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture(scope="module")
def docx(run_pipeline):
    if not os.path.isfile(DOCX_PATH):
        pytest.fail(f"DOCX not found: {DOCX_PATH}")
    return Document(DOCX_PATH)


@pytest.fixture(scope="module")
def interface_tables_in_docx(docx):
    return [
        t for t in docx.tables
        if t.rows and t.rows[0].cells[COL_IF_ID].text.strip() == "Interface ID"
    ]


@pytest.fixture(scope="module")
def all_interface_rows(interface_tables_in_docx):
    return [row for t in interface_tables_in_docx for row in t.rows[1:]]


@pytest.fixture(scope="module")
def all_cell_text(all_interface_rows):
    return {cell.text.strip() for row in all_interface_rows for cell in row.cells}


# ---------------------------------------------------------------------------
# Basic document sanity
# ---------------------------------------------------------------------------

def test_docx_exists(run_pipeline):
    assert os.path.isfile(DOCX_PATH), f"DOCX not found: {DOCX_PATH}"


def test_docx_non_empty(docx):
    text = "\n".join(p.text for p in docx.paragraphs)
    assert len(text.strip()) > 100, "DOCX appears empty or near-empty"


# ---------------------------------------------------------------------------
# Interface tables present
# ---------------------------------------------------------------------------

def test_interface_tables_found_in_docx(interface_tables_in_docx):
    assert len(interface_tables_in_docx) > 0, (
        "No interface tables found in DOCX (expected header 'Interface ID')"
    )


def test_interface_table_has_data_rows(all_interface_rows):
    assert len(all_interface_rows) > 0, "Interface tables have no data rows"


def test_interface_ids_start_with_IF(all_interface_rows):
    for row in all_interface_rows:
        iface_id = row.cells[COL_IF_ID].text.strip()
        assert iface_id.startswith("IF_"), f"Interface ID does not start with IF_: '{iface_id}'"


def test_interface_type_values_valid(all_interface_rows):
    valid = {"Function", "Global Variable"}
    for row in all_interface_rows:
        itype = row.cells[COL_IF_TYPE].text.strip()
        assert itype in valid, f"Unexpected Interface Type in DOCX: '{itype}'"


# ---------------------------------------------------------------------------
# Private items absent, public items present
# ---------------------------------------------------------------------------

def test_private_names_absent_from_docx(all_cell_text):
    for name in PRIVATE_NAMES:
        assert name not in all_cell_text, f"Private name '{name}' found in DOCX interface table"


@pytest.mark.parametrize("name,unit", [
    ("coreAdd",        "Core"), ("coreCompute",    "Core"), ("coreLoopSum",  "Core"),
    ("coreCheck",      "Core"), ("coreSumPoint",   "Core"), ("coreSetResult","Core"),
    ("coreProcess",    "Core"), ("coreOrchestrate","Core"), ("coreSetMode",  "Core"),
    ("coreGetCount",   "Core"), ("libAdd",         "Lib"),  ("libNormalize", "Lib"),
    ("utilCompute",    "Util"), ("utilScale",      "Util"),
    ("g_result",       "Core"), ("g_utilBase",     "Util"),
])
def test_public_name_in_docx(all_cell_text, name, unit):
    assert name in all_cell_text, f"'{name}' ({unit}) missing from DOCX interface table"


# ---------------------------------------------------------------------------
# Direction values in correct column
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("name,expected_direction", [
    ("coreGetCount",  "Out"),
    ("coreSetResult", "In"),
    ("g_result",      "In/Out"),
])
def test_direction_in_docx(all_interface_rows, name, expected_direction):
    row = next((r for r in all_interface_rows if r.cells[COL_IF_NAME].text.strip() == name), None)
    assert row is not None, f"'{name}' row not found in DOCX interface table"
    actual = row.cells[COL_DIRECTION].text.strip()
    assert actual == expected_direction, (
        f"'{name}' direction should be {expected_direction}, got '{actual}'"
    )


# ---------------------------------------------------------------------------
# Images and headings
# ---------------------------------------------------------------------------

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_UNIT_NAMES = ("Core", "Lib", "Util")


def _para_has_image(para):
    """True if this paragraph contains an embedded image (w:drawing element)."""
    return bool(para._element.findall(f".//{{{_W_NS}}}drawing"))


def _collect_unit_diagram_placement(docx):
    """Walk paragraphs to find each unit's level-3 heading, then check whether
    the next paragraph before the following heading contains an image.

    Unit headings are rendered as level-3 with text like '2.1.1 Core'.
    The unit diagram PNG is inserted by doc.add_picture() immediately after,
    before the level-4 '2.1.1.1 unit header' heading.

    Returns dict: unit_name -> True/False.
    """
    paras = list(docx.paragraphs)
    result = {}
    for i, para in enumerate(paras):
        if para.style.name != "Heading 3":
            continue
        text = para.text.strip()
        matched = next((u for u in _UNIT_NAMES if text.endswith(u)), None)
        if matched is None:
            continue
        image_found = False
        for j in range(i + 1, len(paras)):
            if paras[j].style.name.startswith("Heading"):
                break
            if _para_has_image(paras[j]):
                image_found = True
                break
        result[matched] = image_found
    return result


@pytest.fixture(scope="module")
def unit_diagram_placement(docx):
    return _collect_unit_diagram_placement(docx)


def test_docx_has_embedded_images(docx):
    assert len(docx.inline_shapes) > 0, (
        "No embedded images in DOCX — unit/behaviour/flowchart/module-static diagrams missing"
    )


def test_all_unit_headings_found_in_docx(unit_diagram_placement):
    """All three unit-level headings must exist in the DOCX."""
    missing = [u for u in _UNIT_NAMES if u not in unit_diagram_placement]
    assert not missing, f"Unit headings not found in DOCX: {missing}"


@pytest.mark.parametrize("unit", _UNIT_NAMES)
def test_unit_diagram_image_placed_after_heading(unit_diagram_placement, unit):
    """A unit diagram image must appear immediately after each unit's heading,
    before the next sub-heading ('unit header').  Fails if the PNG was missing
    or the exporter placed the image in the wrong section."""
    assert unit_diagram_placement.get(unit), (
        f"No unit diagram image found after '{unit}' heading in DOCX"
    )


@pytest.mark.parametrize("heading", ["Dynamic Behaviour", "Static"])
def test_heading_present(docx, heading):
    headings = {p.text.strip() for p in docx.paragraphs if p.style.name.startswith("Heading")}
    assert any(heading in h for h in headings), (
        f"'{heading}' heading not found. Headings: {sorted(headings)[:10]}"
    )


# ---------------------------------------------------------------------------
# Document structure — sections and headings
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("expected", ["Purpose", "Scope", "Terms"])
def test_introduction_section_headings(docx, expected):
    headings = [p.text.strip() for p in docx.paragraphs if p.style.name.startswith("Heading")]
    assert any(expected in h for h in headings), f"Introduction section '{expected}' not found in DOCX"


@pytest.mark.parametrize("module", ["Core", "Lib", "Util"])
def test_module_level1_heading_present(docx, module):
    level1 = [p.text.strip() for p in docx.paragraphs if p.style.name == "Heading 1"]
    assert any(module in h for h in level1), f"No level-1 heading for module '{module}'"


def test_code_metrics_heading_present(docx):
    headings = {p.text.strip() for p in docx.paragraphs if p.style.name.startswith("Heading")}
    assert any("Code Metrics" in h for h in headings), "Code Metrics heading not found in DOCX"


def test_appendix_heading_present(docx):
    headings = {p.text.strip() for p in docx.paragraphs if p.style.name.startswith("Heading")}
    assert any("Appendix" in h for h in headings), "Appendix heading not found in DOCX"


# ---------------------------------------------------------------------------
# Dynamic Behaviour section
# ---------------------------------------------------------------------------

def test_dynamic_behaviour_sub_headings_for_core(docx):
    """Core is called by external units (App/Main, Cross/Hub) outside the Sample group.
    The DOCX must contain level-3 headings of the form 'N.2.M Core - <func> (<caller>)'."""
    level3 = [p.text.strip() for p in docx.paragraphs if p.style.name == "Heading 3"]
    behaviour_headings = [h for h in level3 if "Core -" in h and "(" in h]
    assert behaviour_headings, (
        "No Dynamic Behaviour sub-headings found for Core "
        "(expected: Core is called by external App/Main and Cross/Hub)"
    )


def test_behaviour_description_tables_present(docx):
    """_add_behavior_description_table creates 5-row tables with 'Requirements','Risk','Capacity'."""
    row_labels = {
        row.cells[0].text.strip()
        for t in docx.tables
        for row in t.rows
    }
    for label in ("Requirements", "Risk", "Capacity", "Input Name", "Output Name"):
        assert label in row_labels, (
            f"Behaviour description table missing '{label}' row label in DOCX"
        )


# ---------------------------------------------------------------------------
# Flowchart tables (Static Design interface section)
# ---------------------------------------------------------------------------

def test_flowchart_tables_present(docx):
    """_add_flowchart_table uses 'Capacity(Density)' — distinct from the behaviour table."""
    row_labels = {
        row.cells[0].text.strip()
        for t in docx.tables
        for row in t.rows
    }
    assert "Capacity(Density)" in row_labels, (
        "No flowchart tables found in DOCX (expected 'Capacity(Density)' row label)"
    )


# ---------------------------------------------------------------------------
# Component/Unit table (_add_component_unit_table)
# ---------------------------------------------------------------------------

def test_component_unit_table_present(docx):
    """_add_component_unit_table writes a 4-column table with headers Component/Unit/Description/Note."""
    all_headers = set()
    for t in docx.tables:
        if t.rows:
            row0_texts = {c.text.strip() for c in t.rows[0].cells}
            all_headers.update(row0_texts)
    for expected in ("Component", "Unit", "Description", "Note"):
        assert expected in all_headers, (
            f"Component/Unit table header '{expected}' not found in DOCX tables"
        )


def test_component_unit_table_has_module_names(docx):
    """Component column must contain the Sample group module names."""
    component_tables = [
        t for t in docx.tables
        if t.rows and {c.text.strip() for c in t.rows[0].cells} >= {"Component", "Unit"}
    ]
    assert component_tables, "No Component/Unit table found in DOCX"
    all_cell_text = {
        c.text.strip()
        for t in component_tables
        for row in t.rows
        for c in row.cells
    }
    for module in ("Core", "Lib", "Util"):
        assert module in all_cell_text, f"Module '{module}' missing from Component/Unit table"


# ---------------------------------------------------------------------------
# Unit header table (_add_unit_header_table)
# ---------------------------------------------------------------------------

def test_unit_header_table_present(docx):
    """_add_unit_header_table writes a 2-column table with 'global variables / typedef / enum / define' header."""
    header_text = "global variables / typedef / enum / define"
    found = any(
        t.rows and t.rows[0].cells[0].text.strip() == header_text
        for t in docx.tables
    )
    assert found, (
        f"Unit header table not found in DOCX (expected first-column header: '{header_text}')"
    )


# ---------------------------------------------------------------------------
# moduleStaticDiagram — PNG or Mermaid text present in Static Design section
# ---------------------------------------------------------------------------

def test_module_static_diagram_content_present(docx):
    """Static Design section must contain either an embedded image (PNG rendered)
    or a paragraph with 'flowchart' (Mermaid text fallback)."""
    paras = list(docx.paragraphs)
    static_headings = [
        i for i, p in enumerate(paras)
        if p.style.name == "Heading 2" and "Static Design" in p.text
    ]
    assert static_headings, "No 'Static Design' Heading 2 found in DOCX"

    _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    for h_idx in static_headings:
        found = False
        for j in range(h_idx + 1, len(paras)):
            if paras[j].style.name.startswith("Heading") and j != h_idx:
                break
            if paras[j]._element.findall(f".//{{{_W_NS}}}drawing"):
                found = True
                break
            if "flowchart" in paras[j].text.lower():
                found = True
                break
        assert found, (
            f"Static Design section (para {h_idx}) has no module diagram image or Mermaid text"
        )
